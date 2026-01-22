"""
Template Creator Service

Processes DOCX and PDF files to:
1. Detect yellow highlighted text sections
2. Replace highlights with placeholder variables ({{field_N}})
3. Generate a clean template file
4. Return a Template State JSON with field metadata
"""

import os
import uuid
import json
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional
from dataclasses import dataclass, asdict
from enum import Enum

import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX


@dataclass
class DetectedField:
    """Represents a detected highlighted field."""
    id: int
    original_text: str
    field_type: str  # "text", "date", "number", etc.
    placeholder: str
    page: Optional[int] = None  # For PDFs
    paragraph: Optional[int] = None  # For DOCX
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "original_text": self.original_text,
            "type": self.field_type,
            "placeholder": self.placeholder,
            "page": self.page,
            "paragraph": self.paragraph,
        }


@dataclass
class TemplateState:
    """The output template state containing template path and fields."""
    template_file_path: str
    original_file_path: str
    fields: List[DetectedField]
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "template_file_path": self.template_file_path,
            "original_file_path": self.original_file_path,
            "fields": [field.to_dict() for field in self.fields],
        }
    
    def to_json(self) -> str:
        return json.dumps(self.to_dict(), indent=2)


# Yellow highlight color indices in Word
# WD_COLOR_INDEX values: YELLOW=7, BRIGHT_GREEN=4
# Also include common numeric values that represent yellow
YELLOW_HIGHLIGHT_INDICES = [7, 11, 4]  # yellow, bright yellow, bright green (often appears yellow)


def process_docx_template(
    file_content: bytes,
    output_dir: str,
    original_filename: str = "document.docx"
) -> TemplateState:
    """
    Process a DOCX file to create a template with placeholders.
    
    Args:
        file_content: The DOCX file content as bytes
        output_dir: Directory to save the template file
        original_filename: Original filename for reference
        
    Returns:
        TemplateState containing template path and detected fields
    """
    # Load document from bytes
    print(f"DEBUG: Loading DOCX document (size: {len(file_content)} bytes)")
    try:
        doc = Document(BytesIO(file_content))
    except Exception as e:
        print(f"DEBUG: Failed to load DOCX: {e}")
        raise ValueError(f"Invalid DOCX file: {e}")
    
    print("DEBUG: DOCX loaded successfully")
    fields: List[DetectedField] = []
    field_counter = 1
    
    # Iterate through all paragraphs
    print(f"DEBUG: Processing {len(doc.paragraphs)} paragraphs")
    for para_idx, paragraph in enumerate(doc.paragraphs):
        try:
            field_counter = _process_paragraph_runs(
                paragraph, para_idx, fields, field_counter
            )
        except Exception as e:
            print(f"DEBUG: Error processing paragraph {para_idx}: {e}")
            raise
    
    # Also process tables
    print(f"DEBUG: Processing {len(doc.tables)} tables")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    field_counter = _process_paragraph_runs(
                        paragraph, para_idx, fields, field_counter
                    )
    
    # Generate unique template filename
    template_id = str(uuid.uuid4())[:8]
    base_name = os.path.splitext(original_filename)[0]
    template_filename = f"{base_name}_template_{template_id}.docx"
    template_path = os.path.join(output_dir, template_filename)
    
    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    # Save the template document
    print(f"DEBUG: Saving template to {template_path}")
    try:
        doc.save(template_path)
    except Exception as e:
        print(f"DEBUG: Failed to save template: {e}")
        raise
    
    print(f"DEBUG: Template saved. Detected {len(fields)} fields.")
    
    return TemplateState(
        template_file_path=template_path,
        original_file_path=original_filename,
        fields=fields,
    )


def process_pdf_template(
    file_content: bytes,
    output_dir: str,
    original_filename: str = "document.pdf"
) -> TemplateState:
    """
    Process a PDF file to create a template with placeholders.
    
    Note: PDF modification is more complex than DOCX. This implementation:
    1. Extracts highlighted text and their positions
    2. Creates annotations or overlays for placeholders
    3. Stores the mapping for later document generation
    
    Args:
        file_content: The PDF file content as bytes
        output_dir: Directory to save the template file
        original_filename: Original filename for reference
        
    Returns:
        TemplateState containing template path and detected fields
    """
    # Open PDF from bytes
    doc = fitz.open(stream=file_content, filetype="pdf")
    
    fields: List[DetectedField] = []
    field_counter = 1
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get all annotations on the page
        annotations = page.annots()
        
        if annotations:
            annots_to_process = list(annotations)
            
            for annot in annots_to_process:
                # Check if it's a highlight annotation (type 8)
                if annot.type[0] == 8:
                    # Get the highlight color
                    colors = annot.colors
                    stroke_color = colors.get("stroke", [1, 1, 0]) if colors else [1, 1, 0]
                    
                    # Check for yellow highlight
                    if _is_yellow_color(stroke_color):
                        # Get text within the highlight rectangle
                        rect = annot.rect
                        highlighted_text = page.get_text("text", clip=rect).strip()
                        
                        if highlighted_text:
                            placeholder = f"{{{{field_{field_counter}}}}}"
                            
                            fields.append(DetectedField(
                                id=field_counter,
                                original_text=highlighted_text,
                                field_type=_infer_field_type(highlighted_text),
                                placeholder=placeholder,
                                page=page_num + 1,
                            ))
                            
                            # Add a text annotation with the placeholder info
                            # This marks where the field should be replaced
                            annot.set_info({
                                "title": f"HighlightEdit Field {field_counter}",
                                "content": json.dumps({
                                    "placeholder": placeholder,
                                    "original": highlighted_text,
                                }),
                            })
                            annot.update()
                            
                            field_counter += 1
        
        # Also check for text with yellow background (drawn highlights)
        # Some PDFs use drawing operations instead of annotations
        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if block.get("type") == 0:  # Text block
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        # Check if span has yellow background color
                        bg_color = span.get("color", 0)
                        # Yellow in PDF color space (approximation)
                        if _is_yellow_pdf_color(bg_color):
                            text = span.get("text", "").strip()
                            if text and not _text_already_captured(text, fields):
                                placeholder = f"{{{{field_{field_counter}}}}}"
                                
                                fields.append(DetectedField(
                                    id=field_counter,
                                    original_text=text,
                                    field_type=_infer_field_type(text),
                                    placeholder=placeholder,
                                    page=page_num + 1,
                                ))
                                
                                field_counter += 1
    
    # Generate unique template filename
    template_id = str(uuid.uuid4())[:8]
    base_name = os.path.splitext(original_filename)[0]
    template_filename = f"{base_name}_template_{template_id}.pdf"
    template_path = os.path.join(output_dir, template_filename)
    
    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    # Save the modified PDF
    doc.save(template_path)
    doc.close()
    
    return TemplateState(
        template_file_path=template_path,
        original_file_path=original_filename,
        fields=fields,
    )


def create_template(
    file_content: bytes,
    filename: str,
    output_dir: str = "./templates"
) -> TemplateState:
    """
    Main entry point: Create a template from a DOCX or PDF file.
    
    Args:
        file_content: The file content as bytes
        filename: Original filename (used to determine file type)
        output_dir: Directory to save template files
        
    Returns:
        TemplateState with template path and detected fields
        
    Raises:
        ValueError: If file type is not supported
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".docx"):
        return process_docx_template(file_content, output_dir, filename)
    elif filename_lower.endswith(".pdf"):
        return process_pdf_template(file_content, output_dir, filename)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Only .docx and .pdf are supported.")


# ============ Helper Functions ============

def _process_paragraph_runs(paragraph, para_idx: int, fields: List[DetectedField], field_counter: int) -> int:
    """
    Process runs in a paragraph, merging adjacent highlighted runs into single fields.
    
    This fixes the issue where Word splits highlighted text like "Party A - Name" into
    multiple runs (e.g., "Party A" and " - Name" as separate runs).
    
    Args:
        paragraph: The docx paragraph object
        para_idx: Paragraph index for reference
        fields: List to append detected fields to
        field_counter: Current field counter
        
    Returns:
        Updated field_counter
    """
    runs = list(paragraph.runs)
    if not runs:
        return field_counter
    
    i = 0
    while i < len(runs):
        run = runs[i]
        
        # Check if this run has yellow highlight
        if not _is_yellow_highlight(run.font.highlight_color):
            i += 1
            continue
        
        # Found a highlighted run - now collect all adjacent highlighted runs
        highlighted_runs = [run]
        combined_text_parts = [run.text]
        
        # Look ahead for more adjacent highlighted runs
        j = i + 1
        while j < len(runs):
            next_run = runs[j]
            if _is_yellow_highlight(next_run.font.highlight_color):
                highlighted_runs.append(next_run)
                combined_text_parts.append(next_run.text)
                j += 1
            else:
                break
        
        # Combine the text from all adjacent highlighted runs
        combined_text = "".join(combined_text_parts).strip()
        
        if combined_text:  # Only process non-empty text
            # Create placeholder
            placeholder = f"{{{{field_{field_counter}}}}}"
            
            # Store the detected field with combined text
            fields.append(DetectedField(
                id=field_counter,
                original_text=combined_text,
                field_type=_infer_field_type(combined_text),
                placeholder=placeholder,
                paragraph=para_idx + 1,
            ))
            
            # Replace the first highlighted run with the placeholder
            highlighted_runs[0].text = placeholder
            highlighted_runs[0].font.highlight_color = None
            
            # Clear the text from subsequent runs (they're now merged)
            for subsequent_run in highlighted_runs[1:]:
                subsequent_run.text = ""
                subsequent_run.font.highlight_color = None
            
            field_counter += 1
        
        # Skip past all the runs we just processed
        i = j
    
    return field_counter


def _is_yellow_highlight(color) -> bool:
    """Check if a Word highlight color is yellow."""
    if color is None:
        return False
    
    # Handle WD_COLOR_INDEX enum values
    try:
        # Direct enum comparison
        if color == WD_COLOR_INDEX.YELLOW:
            return True
        if color == WD_COLOR_INDEX.BRIGHT_GREEN:
            return True
            
        # Check if it has a value attribute (enum)
        if hasattr(color, 'value'):
            return color.value in YELLOW_HIGHLIGHT_INDICES
        
        # Check direct integer value
        if isinstance(color, int):
            return color in YELLOW_HIGHLIGHT_INDICES
            
        # String comparison for some edge cases
        if hasattr(color, 'name'):
            return color.name.upper() in ['YELLOW', 'BRIGHT_GREEN']
            
    except Exception:
        pass
    
    return False


def _is_yellow_color(rgb: List[float], tolerance: float = 0.3) -> bool:
    """Check if an RGB color (0-1 scale) is approximately yellow."""
    if not rgb or len(rgb) < 3:
        return False
    r, g, b = rgb[:3]
    # Yellow = high red, high green, low blue
    return r > 0.7 and g > 0.7 and b < tolerance


def _is_yellow_pdf_color(color_int: int) -> bool:
    """Check if a PDF color integer represents yellow."""
    # Convert int to RGB components
    if color_int == 0:
        return False
    # Common yellow values in PDF: 0xFFFF00, etc.
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return r > 200 and g > 200 and b < 100


def _text_already_captured(text: str, fields: List[DetectedField]) -> bool:
    """Check if text was already captured as a field."""
    return any(f.original_text == text for f in fields)


def _infer_field_type(text: str) -> str:
    """
    Infer the field type based on the original text content.
    
    Returns one of: "text", "date", "email", "phone", "number", "currency"
    """
    text_lower = text.lower().strip()
    
    # Date patterns
    import re
    date_patterns = [
        r'\d{1,2}/\d{1,2}/\d{2,4}',  # MM/DD/YYYY or DD/MM/YYYY
        r'\d{1,2}-\d{1,2}-\d{2,4}',  # MM-DD-YYYY
        r'\d{4}-\d{2}-\d{2}',         # YYYY-MM-DD
        r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)',
    ]
    for pattern in date_patterns:
        if re.search(pattern, text_lower):
            return "date"
    
    # Email pattern
    if re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text):
        return "email"
    
    # Phone pattern
    if re.search(r'[\d\s\-\(\)]{10,}', text):
        return "phone"
    
    # Currency pattern
    if re.search(r'[$€£¥]\s*[\d,]+\.?\d*', text) or re.search(r'[\d,]+\.?\d*\s*[$€£¥]', text):
        return "currency"
    
    # Number pattern (pure numbers)
    if re.match(r'^[\d,]+\.?\d*$', text.strip()):
        return "number"
    
    # Default to text
    return "text"


# ============ Async Wrappers for FastAPI ============

async def create_template_async(
    file_content: bytes,
    filename: str,
    output_dir: str = "./templates"
) -> TemplateState:
    """Async wrapper for create_template."""
    # For CPU-bound work, we could use run_in_executor
    # For now, just call synchronously (document processing is fast)
    return create_template(file_content, filename, output_dir)
