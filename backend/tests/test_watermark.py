"""Tests for the footer watermark implementation."""
import pytest
from docx import Document
from app.services.watermark import add_watermark


def make_simple_doc() -> Document:
    """Create a minimal Document with one paragraph of content."""
    doc = Document()
    doc.add_paragraph("This is the document content.")
    return doc


def test_watermark_adds_footer_paragraph():
    """The last paragraph of the document should contain the watermark text."""
    doc = make_simple_doc()
    add_watermark(doc)
    last_para = doc.paragraphs[-1]
    full_text = "".join(run.text for run in last_para.runs)
    assert "HighlightEdit" in full_text
    assert "Free Plan" in full_text


def test_watermark_does_not_use_header():
    """The document header should NOT contain VML watermark shapes after the rewrite."""
    doc = make_simple_doc()
    add_watermark(doc)
    for section in doc.sections:
        header_xml = section.header._element.xml
        # "PowerPlusWaterMarkObject" is the literal id in the old VML XML template
        assert "PowerPlusWaterMarkObject" not in header_xml, \
            "VML watermark shape found in header — old watermark not removed"


def test_original_content_preserved():
    """The original document content should still be present after watermarking."""
    doc = make_simple_doc()
    add_watermark(doc)
    texts = [p.text for p in doc.paragraphs]
    assert any("document content" in t for t in texts)
