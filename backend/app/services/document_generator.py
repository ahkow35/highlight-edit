"""
Document Generator Service

Generates final documents by:
1. Loading template files with placeholders
2. Injecting user values into placeholders
3. Stripping all highlight formatting
4. Returning clean documents with original font styles preserved
"""

import os
import re
from io import BytesIO
from typing import Dict, Optional
from copy import deepcopy

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn


def generate_docx(
    template_path: str,
    field_values: Dict[str, str],
    output_path: Optional[str] = None,
) -> bytes:
    """
    Generate a DOCX document with placeholders replaced by user values.
    
    Args:
        template_path: Path to the template DOCX file
        field_values: Dict mapping field names (field_1, field_2, etc.) to values
        output_path: Optional path to save the output file
        
    Returns:
        Generated document as bytes
    """
    # Load the template
    doc = Document(template_path)
    
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        _process_paragraph(paragraph, field_values)
    
    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _process_paragraph(paragraph, field_values)
    
    # Process headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                _process_paragraph(paragraph, field_values)
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _process_paragraph(paragraph, field_values)
    
    # Save to bytes
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    # Optionally save to file
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(output_buffer.getvalue())
        output_buffer.seek(0)
    
    return output_buffer.getvalue()


def _process_paragraph(paragraph, field_values: Dict[str, str]):
    """
    Process a paragraph: replace placeholders and strip highlights.
    
    Preserves original font styles while:
    - Replacing {{field_N}} placeholders with user values
    - Removing any remaining yellow highlight formatting
    """
    for run in paragraph.runs:
        text = run.text
        
        # Replace placeholders with user values
        for field_name, value in field_values.items():
            # Support multiple placeholder formats
            patterns = [
                f"{{{{{field_name}}}}}",      # {{field_1}}
                f"{{{{ {field_name} }}}}",    # {{ field_1 }}
            ]
            for pattern in patterns:
                if pattern in text:
                    text = text.replace(pattern, value)
        
        if text != run.text:
            run.text = text
        
        # CRITICAL: Strip ALL highlight formatting
        _strip_highlight(run)


def _strip_highlight(run):
    """
    Remove highlight formatting from a run while preserving other styles.
    
    This ensures the output document has no yellow highlights.
    """
    # Method 1: Set highlight_color to None via python-docx
    try:
        run.font.highlight_color = None
    except Exception:
        pass
    
    # Method 2: Remove highlight via XML (more thorough)
    try:
        # Access the run's XML properties
        rPr = run._r.get_or_add_rPr()
        
        # Find and remove highlight element
        highlight_elem = rPr.find(qn('w:highlight'))
        if highlight_elem is not None:
            rPr.remove(highlight_elem)
        
        # Also check for shading (sometimes used instead of highlight)
        shd_elem = rPr.find(qn('w:shd'))
        if shd_elem is not None:
            # Check if it's a yellow-ish color
            fill = shd_elem.get(qn('w:fill'), '').upper()
            if fill in ['FFFF00', 'FFFF99', 'FFFFCC', 'FFFFE0', 'YELLOW']:
                rPr.remove(shd_elem)
                
    except Exception:
        pass


def generate_docx_stream(
    template_path: str,
    field_values: Dict[str, str],
) -> BytesIO:
    """
    Generate a DOCX document and return as a BytesIO stream.
    
    Args:
        template_path: Path to the template DOCX file
        field_values: Dict mapping field names to values
        
    Returns:
        BytesIO stream containing the generated document
    """
    doc_bytes = generate_docx(template_path, field_values)
    return BytesIO(doc_bytes)


def strip_all_highlights_from_docx(doc_path: str, output_path: str) -> None:
    """
    Utility function to strip all highlights from a DOCX file.
    
    Useful for cleaning up template files or existing documents.
    """
    doc = Document(doc_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _strip_highlight(run)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _strip_highlight(run)
    
    doc.save(output_path)


def get_output_filename(template_path: str, prefix: str = "filled_") -> str:
    """Generate an output filename based on the template path."""
    base_name = os.path.basename(template_path)
    name, ext = os.path.splitext(base_name)
    
    # Remove _template suffix if present
    if name.endswith("_template"):
        name = name[:-9]
    elif "_template_" in name:
        name = name.split("_template_")[0]
    
    return f"{prefix}{name}{ext}"



